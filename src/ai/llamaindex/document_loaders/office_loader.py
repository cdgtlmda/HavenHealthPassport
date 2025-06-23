"""Office document loader for medical documents.

Handles loading of Microsoft Office and similar office suite documents.
"""

import logging
from pathlib import Path
from typing import TYPE_CHECKING, Any, Dict, List, Optional

from llama_index.core import Document

from .base import (
    BaseDocumentLoader,
    DocumentLoaderConfig,
    DocumentMetadata,
    LoaderResult,
)

if TYPE_CHECKING:
    import docx
    import pandas as pd
else:
    try:
        import docx
    except ImportError:
        docx = None  # type: ignore[assignment]

    try:
        import pandas as pd
    except ImportError:
        pd = None  # type: ignore[assignment]

logger = logging.getLogger(__name__)


class OfficeMedicalLoader(BaseDocumentLoader):
    """Loader for office documents (doc, docx, xls, xlsx)."""

    def __init__(self, config: Optional[DocumentLoaderConfig] = None):
        """Initialize office document loader."""
        super().__init__(config or DocumentLoaderConfig())
        self.supported_extensions = [".doc", ".docx", ".xls", ".xlsx"]

    def load(self, file_path: str, **kwargs: Any) -> LoaderResult:
        """Load office document.

        Args:
            file_path: Path to office document

        Returns:
            LoaderResult with documents and metadata
        """
        try:
            path = Path(file_path)

            if not path.exists():
                return LoaderResult(
                    success=False, errors=[f"File not found: {file_path}"]
                )

            # Check file extension
            extension = path.suffix.lower()
            if extension not in self.supported_extensions:
                return LoaderResult(
                    success=False, errors=[f"Unsupported file extension: {extension}"]
                )

            # Load based on file type
            if extension in [".doc", ".docx"]:
                documents = self._load_word_document(path)
            else:  # .xls, .xlsx
                documents = self._load_excel_document(path)

            return LoaderResult(
                success=True,
                documents=documents,
                metadata=DocumentMetadata(
                    file_path=str(path),
                    file_type=extension[1:],
                    file_size=path.stat().st_size,
                ),
            )

        except (OSError, ValueError, IOError, AttributeError) as e:
            logger.error("Error loading office document %s: %s", file_path, str(e))
            return LoaderResult(success=False, errors=[str(e)])

    def _load_word_document(self, path: Path) -> List[Document]:
        """Load Word document."""
        try:
            if docx is None:
                raise ImportError("python-docx is required for Word document loading")

            doc = docx.Document(str(path))

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    table_texts.append("\n".join(table_text))

            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\nTables:\n" + "\n\n".join(table_texts)

            # Extract metadata
            metadata = {
                "source": str(path),
                "file_type": "word",
                "created": (
                    str(doc.core_properties.created)
                    if doc.core_properties.created
                    else None
                ),
                "modified": (
                    str(doc.core_properties.modified)
                    if doc.core_properties.modified
                    else None
                ),
                "author": doc.core_properties.author or "Unknown",
                "title": doc.core_properties.title or path.stem,
                "subject": doc.core_properties.subject,
                "num_paragraphs": len(paragraphs),
                "num_tables": len(doc.tables),
            }

            return [Document(text=full_text, metadata=metadata)]

        except ImportError:
            logger.error(
                "python-docx not installed. Install with: pip install python-docx"
            )
            raise
        except (OSError, ValueError, AttributeError) as e:
            logger.error("Error loading Word document: %s", str(e))
            raise

    def _load_excel_document(self, path: Path) -> List[Document]:
        """Load Excel document."""
        try:
            if pd is None:
                raise ImportError("pandas is required for Excel document loading")

            # Read all sheets
            excel_file = pd.ExcelFile(str(path))
            documents = []

            for sheet_name in excel_file.sheet_names:
                # Read sheet
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                # Convert to text
                text_lines = [f"Sheet: {sheet_name}"]

                # Add column headers
                text_lines.append("Columns: " + ", ".join(df.columns))

                # Add data rows (limit to reasonable number for text processing)
                max_rows = min(len(df), 1000)  # Limit to 1000 rows
                if len(df) > 0:
                    text_lines.append(
                        f"\nData ({len(df)} rows, showing first {max_rows}):"
                    )
                    for _idx, row in df.head(max_rows).iterrows():
                        row_text = " | ".join(
                            [
                                f"{col}: {val}"
                                for col, val in row.items()
                                if pd.notna(val)
                            ]
                        )
                        if row_text:
                            text_lines.append(row_text)

                # Add summary statistics for numeric columns
                numeric_cols = df.select_dtypes(include=["number"]).columns
                if len(numeric_cols) > 0:
                    text_lines.append("\nNumeric Column Statistics:")
                    for col in numeric_cols:
                        stats = df[col].describe()
                        text_lines.append(
                            f"{col}: mean={stats['mean']:.2f}, std={stats['std']:.2f}, min={stats['min']}, max={stats['max']}"
                        )

                # Create document for this sheet
                sheet_text = "\n".join(text_lines)
                metadata = {
                    "source": str(path),
                    "file_type": "excel",
                    "sheet_name": sheet_name,
                    "num_rows": len(df),
                    "num_columns": len(df.columns),
                    "columns": list(df.columns),
                }

                documents.append(Document(text=sheet_text, metadata=metadata))

            return documents

        except ImportError:
            logger.error(
                "pandas and openpyxl not installed. Install with: pip install pandas openpyxl"
            )
            raise
        except (OSError, ValueError, AttributeError, KeyError) as e:
            logger.error("Error loading Excel document: %s", str(e))
            raise

    def validate(self, file_path: str) -> Dict[str, Any]:
        """Validate office document."""
        validation_result: Dict[str, Any] = {
            "valid": True,
            "errors": [],
            "warnings": [],
        }

        try:
            path = Path(file_path)

            # Check file exists
            if not path.exists():
                validation_result["valid"] = False
                validation_result["errors"].append(f"File not found: {file_path}")
                return validation_result

            # Check file extension
            extension = path.suffix.lower()
            if extension not in self.supported_extensions:
                validation_result["valid"] = False
                validation_result["errors"].append(
                    f"Unsupported file type: {extension}"
                )
                return validation_result

            # Check file size
            file_size = path.stat().st_size
            if file_size > self.config.max_file_size_mb * 1024 * 1024:
                validation_result["valid"] = False
                validation_result["errors"].append(
                    f"File too large: {file_size} bytes (max: {self.config.max_file_size_mb * 1024 * 1024})"
                )

            # Try to open the file
            if extension in [".doc", ".docx"]:
                try:
                    if docx is None:
                        raise ImportError(
                            "python-docx is required for Word document validation"
                        )

                    docx.Document(str(path))
                except (ImportError, OSError, ValueError, AttributeError) as e:
                    validation_result["valid"] = False
                    validation_result["errors"].append(f"Invalid Word document: {e}")
            else:  # Excel
                try:
                    if pd is None:
                        raise ImportError(
                            "pandas is required for Excel document validation"
                        )

                    pd.ExcelFile(str(path))
                except (
                    ImportError,
                    OSError,
                    ValueError,
                    AttributeError,
                    KeyError,
                ) as e:
                    validation_result["valid"] = False
                    validation_result["errors"].append(f"Invalid Excel document: {e}")

        except (OSError, ValueError, IOError) as e:
            validation_result["valid"] = False
            validation_result["errors"].append(f"Validation error: {e}")

        return validation_result
